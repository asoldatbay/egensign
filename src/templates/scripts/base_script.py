import base64
import random
import string
from io import BytesIO

from rest_framework.exceptions import APIException, ParseError

from docx import Document
from docx.shared import Pt


class BaseScript:

    template = None
    content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    @classmethod
    def _set_font(cls, doc, font_name, size=12, color=None):
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(size)
                if color:
                    run.font.color.rgb = color

    @classmethod
    def _spacing(cls, doc, number):
        for _ in range(number):
            doc.add_paragraph()

    @classmethod
    def _get_buffer(cls, doc):
        # Save the document to a BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    @classmethod
    def _get_default_filename(cls):
        return "default_filename"

    @classmethod
    def _get_filename(cls, length=64):
        random_string = "".join(
            random.choices(
                string.ascii_uppercase + string.ascii_lowercase + string.digits,
                k=length,
            )
        )
        return f"{cls._get_default_filename()}_{random_string}.docx"

    @classmethod
    def _user_preview_permission(cls, request):
        raise APIException(detail="Implement _user_preview_permission")

    @classmethod
    def _user_upload_permission(cls, request):
        raise APIException(detail="Implement _user_has_permission")

    # @classmethod
    # def _get_bucket_name(cls):
    #     raise APIException(detail="Implement _get_bucket_name")

    # @classmethod
    # def _get_serializer(cls):
    #     raise APIException(detail="Implement _get_serializer")

    @classmethod
    def validate_data(cls, request):
        # serializer_class = cls._get_serializer()
        # serializer: Serializer = serializer_class(data=request.data)
        # serializer.is_valid(raise_exception=True)

        if not request.data.get('filled_data', None):
            raise ParseError(detail='Missing filling data')

        return request.data["filled_data"]

    @classmethod
    def generate_header(cls, doc, data):
        raise APIException(detail="Implement generate_header")

    @classmethod
    def generate_body(cls, doc, data):
        raise APIException(detail="Implement generate_body")

    @classmethod
    def generate_footer(cls, doc, data):
        raise APIException(detail="Implement generate_footer")

    @classmethod
    def generate(cls, data):
        doc = Document()

        cls.generate_header(doc, data)
        cls.generate_body(doc, data)
        cls.generate_footer(doc, data)

        return doc

    def _get_response(self, buffer, filename):
        # Return the document as a response
        file = buffer.getvalue()
        file64 = base64.b64encode(file)
        file_type = self.content_type

        # return HttpResponse(file64, content_type=file_type)
        return file64, file_type

    def preview(self, request):
        self._user_preview_permission(request)

        buffer = self._get_buffer(self.generate(self.validate_data(request)))

        return self._get_response(buffer, "preview.docx")

    # @action(detail=False, methods=['post'])
    # def upload(self, request):
    #     self._user_upload_permission(request)
    #
    #     buffer = self._get_buffer(self.generate(self.validate_data(request)))
    #
    #     bucket = settings.MINIO_BUCKET
    #
    #     client = Minio(
    #         f"{settings.MINIO_HOST}:{settings.MINIO_PORT}",
    #         access_key=settings.MINIO_ACCESS_KEY,
    #         secret_key=settings.MINIO_SECRET_KEY,
    #         secure=settings.MINIO_IS_SECURED,
    #     )
    #
    #     filename = self._get_filename()
    #
    #     client.put_object(bucket, filename, io.BytesIO(buffer.getvalue()), length=len(buffer.getvalue()),
    #                       content_type=self.content_type)
    #
    #     document = DocumentUnit(
    #         user_id=request.user.id,
    #         company_id=request.user.company_id,
    #         url=f"{settings.MINIO_HOST}:{settings.MINIO_PORT}/{settings.MINIO_BUCKET}/{filename}"
    #     )
    #     document.save()
    #
    #     return self._get_response(buffer, filename)

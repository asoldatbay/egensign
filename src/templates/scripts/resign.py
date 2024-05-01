from django.conf import settings
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

from src.templates.models import Template
from src.templates.scripts.base_script import BaseScript


class ResignScript(BaseScript):

    template = Template.objects.get(script_code="script_code")

    @classmethod
    def _get_default_filename(cls):
        return "resign"

    @classmethod
    def _get_bucket_name(cls):
        return settings.MINIO_EMPLOYEES_BUCKET

    @classmethod
    def _user_preview_permission(cls, request):
        pass

    @classmethod
    def _user_upload_permission(cls, request):
        pass

    @classmethod
    def generate_header(cls, doc, data):
        # Add header content
        header_content = f"{data['receiver_position']}\n{data['receiver_fullname']}\n" \
                         f"от сотрудника\n{data['sender_fullname']}"
        header_paragraph = doc.add_paragraph(header_content)
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Align to the right
        cls._set_font(doc, "Calibri")

        cls._spacing(doc, 3)

    @classmethod
    def generate_body(cls, doc, data):
        # Add body title centered and in bold
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run("ЗАЯВЛЕНИЕ")
        title_run.bold = True
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cls._set_font(doc, "Calibri")

        # Add body content with indentation
        body_content = f"\tПрошу уволить меня по собственному желанию с {data['from_date']}"
        body_content += f" {data['reason']}." if data['reason'] else "."
        body_paragraph = doc.add_paragraph(body_content)
        body_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        cls._set_font(doc, "Calibri")

        cls._spacing(doc, 3)

    @classmethod
    def generate_footer(cls, doc, data=None):
        # Add footer with signing line and current date
        footer_paragraph = doc.add_paragraph()
        signing_line = footer_paragraph.add_run("__________________________\nПодпись\n")
        signing_line.add_text(datetime.now().strftime("%d.%m.%Y"))  # Add current date
        signing_line.font.size = Pt(10)
        cls._set_font(doc, "Calibri")
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

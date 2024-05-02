from django.conf import settings
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

from templates.models import Template
from templates.scripts.base_script import BaseScript


class StudentPracticeScript(BaseScript):

    template = Template.objects.get(script_code="fec2d982-8922-42bc-ac87-ea1a6c2ecea3")

    @classmethod
    def _get_default_filename(cls):
        return "student_practice"

    # @classmethod
    # def _get_bucket_name(cls):
    #     return settings.MINIO_EMPLOYEES_BUCKET

    @classmethod
    def _user_preview_permission(cls, request):
        pass

    @classmethod
    def _user_upload_permission(cls, request):
        pass

    @classmethod
    def generate_header(cls, doc, data):
        # Add header content
        header_content = f"Заведующей кафедрой\n" \
                         f"{data['faculty_name']}\n" \
                         f"А.О МУИТ\n" \
                         f"{data['faculty_head']}"
        header_paragraph = doc.add_paragraph(header_content)
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Align to the right
        cls._set_font(doc, "Calibri")

        cls._spacing(doc, 3)

    @classmethod
    def generate_body(cls, doc, data):
        # Add body title centered and in bold
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(
            f"Уважаемая {data['faculty_head_fullname']}"
        )
        title_run.bold = True
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cls._set_font(doc, "Calibri")

        # Add body content with indentation
        '''
            Предприятие TOO PRIME SOURCE не возражает по поводу прохождения преддипломной практики студента 4 курса
            очной формы обучения образовательной программы 6В06106 –
            Вычислительная техника и программное обеспечение Международного университета информационных технологий
            Али Манашов на период с 18 марта по 8 апреля 2024г. На безвозмездной основе. 
        '''
        body_content = f"\t Предприятие {data['student_company']} не возражает по " \
                       f"поводу прохождения преддипломной практики студента " \
                       f"{data['student_course_number']} курса " \
                       f"{data['student_edu_form']} " \
                       f"формы обучения образовательной программы " \
                       f"{data['student_program_code']} - " \
                       f"{data['student_program']} " \
                       f"Международного университета информационных технологий " \
                       f"{data['student_fullname']} на период с" \
                       f"{data['date_from']} по {data['date_to']}. " \
                       f"На безвозмездной основе."
        body_paragraph = doc.add_paragraph(body_content)
        body_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        cls._set_font(doc, "Calibri")

        cls._spacing(doc, 3)

    @classmethod
    def generate_footer(cls, doc, data=None):
        # Add footer with signing line and current date
        footer_paragraph = doc.add_paragraph()
        signing_line = footer_paragraph.add_run("Подпись   ___________\nДата ")
        signing_line.add_text(datetime.now().strftime("%d.%m.%Y"))  # Add current date
        signing_line.font.size = Pt(10)
        cls._set_font(doc, "Calibri")
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
